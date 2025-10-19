# updated_generate_and_push.py
import json
import urllib3
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import requests
import matplotlib.pyplot as plt
import networkx as nx
import os
import base64
from docx.shared import Pt

# --- Load config ---
with open("config_file.json", "r") as f:
    config = json.load(f)

GEMINI_API_URL = config.get("gemini_api_url")
GEMINI_API_KEY = config.get("gemini_api_key")
XML_PATH = config.get("source_xml_path")
DOCX_PATH = config.get("target_docx_path", "output_doc.docx")
GROOVY_SCRIPTS_FOLDER = config.get("groovy_scripts_folder", None)

# GitHub push config
GITHUB_REPO = config.get("github_repo")          # format: owner/repo
GITHUB_BRANCH = config.get("github_branch", "main")
GITHUB_TOKEN = config.get("github_token")        # personal access token (repo scope)
REPO_TARGET_PATH = config.get("repo_target_path", "docs")  # folder in repo to put doc

iflow_name = os.path.splitext(os.path.basename(XML_PATH))[0]

# --- Gemini API Call ---
def call_gemini(prompt):
    if not GEMINI_API_URL or not GEMINI_API_KEY:
        return ""
    headers = {"Content-Type": "application/json", "X-goog-api-key": GEMINI_API_KEY}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        response = requests.post(GEMINI_API_URL, headers=headers, json=data, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result["candidates"][0]["content"]["parts"][0]["text"]
        else:
            print("Gemini error:", response.status_code, response.text)
            return ""
    except Exception as e:
        print("Gemini request failed:", e)
        return ""

# --- Utility Functions ---
def format_key(key):
    import re
    key = key.replace("_", " ")
    key = re.sub(r"(?<!^)(?=[A-Z])", " ", key)
    return key.title()

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, center=False, font_size=12):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)

def add_colored_table(doc, data, column_names, header_color=RGBColor(0, 51, 102)):
    table = doc.add_table(rows=1, cols=len(column_names))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(column_names):
        run = hdr_cells[i].paragraphs[0].add_run(name)
        run.bold = True
        run.font.color.rgb = header_color
    for row in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    return table

def add_header_footer(doc, page_name, author, version):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Technical Specification: iFlow - {iflow_name} "
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer
    footer_para = footer.paragraphs[0]
    today = datetime.today().strftime("%Y-%m-%d")
    footer_para.text = f"Author: {author} | Version: {version} | Date: {today} "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- XML Extraction Functions ---
def extract_section_xml(root, xpath):
    elems = root.findall(xpath)
    return "\n".join([ET.tostring(e, encoding="unicode") for e in elems])

def extract_properties_from_extension(elem):
    props = []
    for ext in elem.findall(".//{http:///com.sap.ifl.model/Ifl.xsd}property"):
        key = ext.findtext("key")
        value = ext.findtext("value")
        if key:
            props.append([format_key(key), value if value else ""])
    return props

def extract_security(root):
    for collab in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}collaboration"):
        ext_elems = collab.findall("{http://www.omg.org/spec/BPMN/20100524/MODEL}extensionElements")
        if ext_elems:
            return extract_properties_from_extension(ext_elems[0])
    return []

def build_id_name_map(root):
    id_name = {}
    for elem in root.iter():
        id_ = elem.attrib.get("id")
        name = elem.attrib.get("name")
        if id_:
            id_name[id_] = name if name else id_
    return id_name

def extract_message_flows_with_names(root, id_name):
    flows = []
    for msg in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}messageFlow"):
        source = msg.attrib.get("sourceRef")
        target = msg.attrib.get("targetRef")
        source_name = id_name.get(source, source)
        target_name = id_name.get(target, target)
        name = msg.attrib.get("name", "")
        flows.append((source_name, target_name, name))
    return flows

def extract_sequence_flows_with_names(root, id_name):
    flows = []
    for seq in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}sequenceFlow"):
        source = seq.attrib.get("sourceRef")
        target = seq.attrib.get("targetRef")
        source_name = id_name.get(source, source)
        target_name = id_name.get(target, target)
        name = seq.attrib.get("name", "")
        flows.append((source_name, target_name, name))
    return flows

def extract_components(root):
    components = []
    for proc in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process"):
        proc_name = proc.attrib.get("name", "Unknown")
        for ext_elem in proc.findall("{http://www.omg.org/spec/BPMN/20100524/MODEL}extensionElements"):
            props = extract_properties_from_extension(ext_elem)
            for key, value in props:
                components.append([proc_name, key, value])
        if not proc.findall("{http://www.omg.org/spec/BPMN/20100524/MODEL}extensionElements"):
            components.append([proc_name, "", ""])
    return components

def extract_components_from_process(proc):
    components = []
    proc_name = proc.attrib.get("name", "Unknown")
    for ext_elem in proc.findall("{http://www.omg.org/spec/BPMN/20100524/MODEL}extensionElements"):
        props = extract_properties_from_extension(ext_elem)
        for key, value in props:
            components.append([proc_name, key, value])
    if not proc.findall("{http://www.omg.org/spec/BPMN/20100524/MODEL}extensionElements"):
        components.append([proc_name, "", ""])
    return components

def extract_child_properties(process_elem):
    results = []
    for child in list(process_elem):
        tag_name = child.tag.split("}")[-1]
        child_name = child.attrib.get("name", "")
        heading = f"{tag_name} {child_name}".strip()
        props = []
        for ext_elem in child:
            if ext_elem.tag.endswith("extensionElements"):
                for prop in ext_elem:
                    if prop.tag.endswith("property"):
                        key = prop.findtext("key")
                        value = prop.findtext("value")
                        if key:
                            props.append((format_key(key), value if value else ""))
        results.append({"heading": heading, "properties": props})
    return results

def extract_sender_properties(root):
    sender_props = []
    ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL", "ifl": "http:///com.sap.ifl.model/Ifl.xsd"}
    for msg in root.findall(".//bpmn2:messageFlow", ns):
        for ext_elem in msg.findall("bpmn2:extensionElements", ns):
            is_sender = False
            for prop in ext_elem.findall("ifl:property", ns):
                key = prop.findtext("key")
                value = prop.findtext("value")
                if key and key.strip().lower() == "direction" and value and value.strip().lower() == "sender":
                    is_sender = True
                    break
            if is_sender:
                for prop in ext_elem.findall("ifl:property", ns):
                    key = prop.findtext("key")
                    value = prop.findtext("value")
                    if key:
                        sender_props.append([format_key(key), value if value else ""])
    return sender_props

def extract_receiver_properties(root):
    receiver_props = []
    ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL", "ifl": "http:///com.sap.ifl.model/Ifl.xsd"}
    for msg in root.findall(".//bpmn2:messageFlow", ns):
        for ext_elem in msg.findall("bpmn2:extensionElements", ns):
            is_receiver = False
            for prop in ext_elem.findall("ifl:property", ns):
                key = prop.findtext("key")
                value = prop.findtext("value")
                if key and key.strip().lower() == "direction" and value and value.strip().lower() == "receiver":
                    is_receiver = True
                    break
            if is_receiver:
                for prop in ext_elem.findall("ifl:property", ns):
                    key = prop.findtext("key")
                    value = prop.findtext("value")
                    if key:
                        receiver_props.append([format_key(key), value if value else ""])
    return receiver_props

def extract_mapping_properties(root):
    ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL", "ifl": "http:///com.sap.ifl.model/Ifl.xsd"}
    mappings = []
    for process in root.findall(".//bpmn2:process", ns):
        for call_activity in process.findall("bpmn2:callActivity", ns):
            for ext_elem in call_activity.findall("bpmn2:extensionElements", ns):
                for prop in ext_elem.findall("ifl:property", ns):
                    key = prop.findtext("key")
                    value = prop.findtext("value")
                    if key and value and key.strip() == "activityType" and value.strip() == "Mapping":
                        mapping_props = []
                        for p in ext_elem.findall("ifl:property", ns):
                            k = p.findtext("key")
                            v = p.findtext("value")
                            if k:
                                mapping_props.append([k, v if v else ""])
                        mappings.append(mapping_props)
    return mappings

def extract_exception_properties(root):
    ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL", "ifl": "http:///com.sap.ifl.model/Ifl.xsd"}
    exceptions = []
    for process in root.findall(".//bpmn2:process", ns):
        for sub_proc in process.findall("bpmn2:subProcess", ns):
            found = False
            for ext_elem in sub_proc.findall("bpmn2:extensionElements", ns):
                for prop in ext_elem.findall("ifl:property", ns):
                    key = prop.findtext("key")
                    value = prop.findtext("value")
                    if key and value and key.strip() == "activityType" and value.strip() == "ErrorEventSubProcessTemplate":
                        found = True
                        break
                if found:
                    subproc_props = []
                    for p in ext_elem.findall("ifl:property", ns):
                        k = p.findtext("key")
                        v = p.findtext("value")
                        if k:
                            subproc_props.append([k, v if v else ""])
                    exceptions.append({"subproc_props": subproc_props, "children": []})
                    for child in list(sub_proc):
                        for ext_elem_child in child.findall("bpmn2:extensionElements", ns):
                            child_props = []
                            for prop in ext_elem_child.findall("ifl:property", ns):
                                k = prop.findtext("key")
                                v = prop.findtext("value")
                                if k:
                                    child_props.append([k, v if v else ""])
                            if child_props:
                                exceptions[-1]["children"].append({"tag": child.tag.split("}")[-1], "name": child.attrib.get("name", ""), "props": child_props})
                    break
    return exceptions

def exception_props_to_xml(exceptions):
    xml = "<Exceptions>\n"
    for idx, exc in enumerate(exceptions, 1):
        xml += f'  <ExceptionSubProcess id="{idx}">\n'
        xml += "    <Properties>\n"
        for key, value in exc["subproc_props"]:
            xml += f"      <Property>\n        <Key>{key}</Key>\n        <Value>{value}</Value>\n      </Property>\n"
        xml += "    </Properties>\n"
        for child in exc["children"]:
            xml += f'    <ChildElement type="{child["tag"]}" name="{child["name"]}">\n'
            for key, value in child["props"]:
                xml += f"      <Property>\n        <Key>{key}</Key>\n        <Value>{value}</Value>\n      </Property>\n"
            xml += "    </ChildElement>\n"
        xml += "  </ExceptionSubProcess>\n"
    xml += "</Exceptions>"
    return xml

def gemini_section_summary(section_name, xml_fragment, extra_context=""):
    prompt = (f"Summarize the following XML for the {section_name} section of an SAP iFlow in a human-friendly, technical style. {extra_context}\nXML:\n{xml_fragment}\nLimit to 5 sentences.")
    return call_gemini(prompt)

def create_bpmn_diagram_horizontal(root, image_path):
    id_name = build_id_name_map(root)
    flows = []
    for seq in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}sequenceFlow"):
        src = seq.attrib.get("sourceRef")
        tgt = seq.attrib.get("targetRef")
        label = seq.attrib.get("name", "")
        flows.append((id_name.get(src, src), id_name.get(tgt, tgt), label))
    nodes = set()
    for src, tgt, _ in flows:
        nodes.add(src)
        nodes.add(tgt)
    G = nx.DiGraph()
    for node in nodes:
        G.add_node(node, shape="box")
    for src, tgt, label in flows:
        G.add_edge(src, tgt, label=label if label else "Sequence")
    try:
        pos = nx.nx_agraph.graphviz_layout(G, prog="dot", args="-Grankdir=LR")
    except Exception:
        pos = nx.spring_layout(G)
    plt.figure(figsize=(12, 4))
    nx.draw(G, pos, with_labels=True, node_color="lightblue", node_size=2000, font_size=10, arrows=True)
    edge_labels = nx.get_edge_attributes(G, "label")
    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_color="red")
    plt.title("BPMN Diagram (Business Names, Left-to-Right)")
    plt.tight_layout()
    plt.savefig(image_path)
    plt.close()

def get_all_groovy_scripts(folder_path):
    scripts = []
    if folder_path and os.path.isdir(folder_path):
        for fname in sorted(os.listdir(folder_path)):
            if fname.lower().endswith(".groovy"):
                fpath = os.path.join(folder_path, fname)
                with open(fpath, "r", encoding="utf-8") as f:
                    scripts.append((fname, f.read()))
    return scripts

# --- Upload to GitHub (Contents API) ---
def upload_to_github(local_file_path, repo, branch="main", repo_folder="docs", token=None):
    """
    Uploads local_file_path to GitHub repo in folder repo_folder (creates/updates).
    repo: "owner/repo"
    """
    if not token:
        print("No GitHub token provided - skipping upload.")
        return {"ok": False, "reason": "no_token"}

    # Build target path e.g., docs/output_doc.docx
    filename = os.path.basename(local_file_path)
    # ensure folder doesn't start or end with '/'
    repo_folder = repo_folder.strip("/")
    if repo_folder == "":
        target_path = filename
    else:
        target_path = f"{repo_folder}/{filename}"

    api_base = "https://api.github.com"
    get_url = f"{api_base}/repos/{repo}/contents/{target_path}"
    headers = {"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}

    # get existing file to obtain sha (if any)
    sha = None
    try:
        resp = requests.get(get_url, headers=headers, params={"ref": branch}, timeout=30)
        if resp.status_code == 200:
            sha = resp.json().get("sha")
    except Exception as e:
        print("Warning: could not check existing file on GitHub:", e)

    # read file and base64 encode
    with open(local_file_path, "rb") as f:
        content_bytes = f.read()
    content_b64 = base64.b64encode(content_bytes).decode("utf-8")

    put_url = get_url
    message = f"Automated: add/update iFlow spec {filename} (generated {datetime.today().strftime('%Y-%m-%d')})"
    payload = {"message": message, "content": content_b64, "branch": branch}
    if sha:
        payload["sha"] = sha

    try:
        put_resp = requests.put(put_url, headers=headers, json=payload, timeout=60)
        if put_resp.status_code in (200, 201):
            print(f"✅ Uploaded {filename} to {repo}/{target_path} on branch {branch}")
            return {"ok": True, "url": put_resp.json().get("content", {}).get("html_url")}
        else:
            print("GitHub upload failed:", put_resp.status_code, put_resp.text)
            return {"ok": False, "status_code": put_resp.status_code, "text": put_resp.text}
    except Exception as e:
        print("GitHub upload exception:", e)
        return {"ok": False, "reason": str(e)}

# --- Main Document Generation ---
def generate_iflow_spec():
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    # Download the iFlow XML if URL
    if XML_PATH.startswith("http://") or XML_PATH.startswith("https://"):
        print(f"Downloading iFlow XML from: {XML_PATH}")
        try:
            response = requests.get(XML_PATH, verify=False, timeout=30)
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            print("Error downloading the iFlow XML:", e)
            return
        local_file = os.path.basename(XML_PATH)
        with open(local_file, "wb") as f:
            f.write(response.content)
        print(f"Downloaded file saved locally as: {local_file}")
        tree = ET.parse(local_file)
        root = tree.getroot()
    else:
        tree = ET.parse(XML_PATH)
        root = tree.getroot()

    doc = Document()
    add_header_footer(doc, f"{iflow_name}Technical Specification", "Generated by AI", "1.0")

    # Title Page
    add_paragraph(doc, "SAP Integration Suite \nCloud Integration - Technical Specification\n " f"iFlow Name : {iflow_name}", bold=True, center=True, font_size=24)
    add_paragraph(doc, "Version: 1.0", center=True)
    add_paragraph(doc, "Author: Generated by AI", center=True)
    add_paragraph(doc, f"Date: {datetime.today().strftime('%Y-%m-%d')}", center=True)
    doc.add_page_break()

    # Table of Contents
    add_heading(doc, "Table of Contents", level=1)
    toc_sections = [
        "1. Change History",
        "2. Overview",
        "3. High level iFlow Design",
        "4. Message Flow",
        "5. Technical Description",
        "5.1. Main Integration Process",
        "5.2. Local Integration Process",
        "5.3. Sender",
        "5.4. Receiver",
        "5.5. Mappings",
        "5.6. Security",
        "5.7. Groovy Scripts",
        "5.8. Error Handling & Logging",
        "6. Version and Metadata",
        "7. Appendix",
    ]
    for section in toc_sections:
        add_paragraph(doc, section)
    doc.add_page_break()

    # 1. Change History
    add_heading(doc, "1. Change History", level=1)
    add_colored_table(doc, [["1.0", datetime.today().strftime("%Y-%m-%d"), "Generated by AI", "Initial version"]], ["Version", "Date", "Author", "Description"])
    doc.add_page_break()

    # 2. Overview
    add_heading(doc, "2. Overview", level=1)
    overview_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}collaboration")
    overview_text = gemini_section_summary("Overview", overview_xml, f"Describe the purpose of this technical specification document for the iFlow named {iflow_name}. Don't explain the iflow but just provide what is the use of this technical specification document .")
    add_paragraph(doc, overview_text)

    # 3. High level iFlow Design
    add_heading(doc, "3. High level iFlow Design", level=1)
    design_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process")
    design_text = gemini_section_summary("High level iFlow Design", design_xml, f"Describe the main components and flow of the message in the iFlow named {iflow_name} from Sender system to Receiver system.")
    add_paragraph(doc, design_text)
    diagram_path = "bpmn_diagram.png"
    create_bpmn_diagram_horizontal(root, diagram_path)
    try:
        doc.add_picture(diagram_path, width=Inches(6))
        add_paragraph(doc, "Figure: High level BPMN iFlow message and sequence flow", center=True)
    except Exception as e:
        print("Could not add diagram image to doc:", e)

    # 4. Message Flow
    add_heading(doc, "4. Message Flow", level=1)
    id_name = build_id_name_map(root)
    message_flows = extract_message_flows_with_names(root, id_name)
    message_flows_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}messageFlow")
    message_flows_text = gemini_section_summary("Message Flow", message_flows_xml)
    add_paragraph(doc, message_flows_text)
    if message_flows:
        add_colored_table(doc, [[src, tgt, label] for src, tgt, label in message_flows], ["Source", "Target", "Name"])

    # 5. Technical Description
    add_heading(doc, "5. Technical Description", level=1)

    # 5.1. Main Integration Process
    add_heading(doc, "5.1. Main Integration Process ", level=2)
    def extract_process_1_xml(root):
        ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL"}
        for process in root.findall(".//bpmn2:process", ns):
            if process.attrib.get("id") == "Process_1":
                return ET.tostring(process, encoding="unicode")
        return ""
    process_1_xml = extract_process_1_xml(root)
    process_1_text = gemini_section_summary("Main Integration Process", process_1_xml, "Summarize the main integration process and its child elements for SAP iFlow Process_1.")
    add_paragraph(doc, process_1_text)

    # components for Process_1
    target_process = None
    for process in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process"):
        if process.attrib.get("id") == "Process_1":
            target_process = process
            break
    components = []
    if target_process is not None:
        components = extract_components_from_process(target_process)
    if components:
        add_colored_table(doc, components, ["Component Name", "Key", "Value"])

    ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL", "ifl": "http:///com.sap.ifl.model/Ifl.xsd"}
    if target_process is not None:
        child_props = extract_child_properties(target_process)
        for item in child_props:
            if item["properties"]:
                add_heading(doc, f"{item['heading']} Properties", level=3)
                add_colored_table(doc, item["properties"], ["Key", "Value"])
    else:
        add_paragraph(doc, "No process with id='Process_1' found.")

    # 5.2. Local Integration Process
    add_heading(doc, "5.2. Local Integration Process ", level=2)
    def extract_process_local_xml(root):
        ns = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL"}
        for process in root.findall(".//bpmn2:process", ns):
            if process.attrib.get("id") != "Process_1":
                return ET.tostring(process, encoding="unicode")
        return ""
    local_process_xml = extract_process_local_xml(root)
    local_process_text = gemini_section_summary("Main Integration Process", local_process_xml, "Summarize the main integration process and its child elements for SAP iFlow Process_1.")
    add_paragraph(doc, local_process_text)

    target_process = None
    for process in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process"):
        if process.attrib.get("id") != "Process_1":
            target_process = process
            break
    components = []
    if target_process is not None:
        components = extract_components_from_process(target_process)
    if components:
        add_colored_table(doc, components, ["Component Name", "Key", "Value"])

    local_process = None
    for process in root.findall(".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process"):
        if process.attrib.get("id") != "Process_1":
            local_process = process
            break
    if local_process is not None:
        child_props = extract_child_properties(local_process)
        for item in child_props:
            if item["properties"]:
                add_heading(doc, f"{item['heading']} Properties", level=3)
                add_colored_table(doc, item["properties"], ["Key", "Value"])
    else:
        add_paragraph(doc, "No process with id='Process_1' found.")

    # 5.3. Sender
    add_heading(doc, "5.3. Sender", level=2)
    def sender_props_to_xml(sender_props):
        xml = "<SenderProperties>\n"
        for key, value in sender_props:
            xml += f"  <Property>\n    <Key>{key}</Key>\n    <Value>{value}</Value>\n  </Property>\n"
        xml += "</SenderProperties>"
        return xml
    sender_props = extract_sender_properties(root)
    sender_props_xml = sender_props_to_xml(sender_props)
    sender_text = gemini_section_summary("Sender", sender_props_xml, "Identify the sender system, protocol, authentication method, and key configuration parameters. Explain the business role of this endpoint.")
    add_paragraph(doc, sender_text)
    if sender_props:
        add_colored_table(doc, sender_props, ["Key", "Value"])

    # 5.4. Receiver
    add_heading(doc, "5.4. Receiver", level=2)
    def receiver_props_to_xml(receiver_props):
        xml = "<ReceiverProperties>\n"
        for key, value in receiver_props:
            xml += f"  <Property>\n    <Key>{key}</Key>\n    <Value>{value}</Value>\n  </Property>\n"
        xml += "</ReceiverProperties>"
        return xml
    receiver_props = extract_receiver_properties(root)
    receiver_props_xml = receiver_props_to_xml(receiver_props)
    receiver_text = gemini_section_summary("Receiver", receiver_props_xml, "Identify receiver components and describe their role.")
    add_paragraph(doc, receiver_text)
    if receiver_props:
        add_colored_table(doc, receiver_props, ["Key", "Value"])

    # 5.5. Mappings
    add_heading(doc, "5.5. Mappings", level=2)
    def mapping_props_to_xml(mapping_props_list):
        xml = "<Mappings>\n"
        for idx, mapping_props in enumerate(mapping_props_list, 1):
            xml += f'  <MappingActivity id="{idx}">\n'
            for key, value in mapping_props:
                xml += f"    <Property>\n      <Key>{key}</Key>\n      <Value>{value}</Value>\n    </Property>\n"
            xml += "  </MappingActivity>\n"
        xml += "</Mappings>"
        return xml
    mapping_props_list = extract_mapping_properties(root)
    mapping_props_xml = mapping_props_to_xml(mapping_props_list)
    mapping_text = gemini_section_summary("Mappings", mapping_props_xml, "Describe any data mapping or transformation logic.")
    add_paragraph(doc, mapping_text)
    mapping_props_list = extract_mapping_properties(root)
    if mapping_props_list:
        for idx, mapping_props in enumerate(mapping_props_list, 1):
            add_heading(doc, f"Mapping Activity {idx} Properties", level=3)
            add_colored_table(doc, mapping_props, ["Key", "Value"])
    else:
        add_paragraph(doc, "No mapping activities found in the iFlow.")

    # 5.6. Security
    add_heading(doc, "5.6. Security", level=2)
    security = extract_security(root)
    security_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}collaboration")
    security_text = gemini_section_summary("Security", security_xml)
    add_paragraph(doc, security_text)
    if security:
        add_colored_table(doc, security, ["Key", "Value"])
    else:
        add_paragraph(doc, "No security properties found.")

    # 5.7. Groovy Scripts
    add_heading(doc, "5.7. Groovy Scripts", level=2)
    components_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process")
    message_handling_text = gemini_section_summary("Groovy Scripts", components_xml, f"Describe how and where groovy script is used in the iflow, {iflow_name}.")
    add_paragraph(doc, message_handling_text)
    groovy_scripts = get_all_groovy_scripts(GROOVY_SCRIPTS_FOLDER)
    if groovy_scripts:
        for fname, content in groovy_scripts:
            add_heading(doc, f"Script: {fname}", level=3)
            explanation = gemini_section_summary(f"Groovy Script: {fname}", content, "Explain in detail what this Groovy script does in the context of SAP Integration Suite iFlow.")
            add_paragraph(doc, explanation)
            p = doc.add_paragraph()
            run = p.add_run(content)
            font = run.font
            font.name = "Courier New"
            font.size = Pt(10)
    else:
        add_paragraph(doc, "No Groovy scripts found in the specified folder.")

    # 5.8. Error Handling & Logging
    add_heading(doc, "5.8. Error Handling & Logging", level=2)
    exceptions = extract_exception_properties(root)
    exceptions_xml = exception_props_to_xml(exceptions)
    error_handling_text = gemini_section_summary("Error Handling & Logging", exceptions_xml, "Describe error handling and logging mechanisms.")
    add_paragraph(doc, error_handling_text)
    if exceptions:
        for idx, exc in enumerate(exceptions, 1):
            add_heading(doc, f"Exception SubProcess {idx} Properties", level=3)
            add_colored_table(doc, exc["subproc_props"], ["Key", "Value"])
            for child in exc["children"]:
                add_heading(doc, f'Child Element: {child["tag"]} {child["name"]}', level=4)
                add_colored_table(doc, child["props"], ["Key", "Value"])
    else:
        add_paragraph(doc, "No exception subprocesses found in the iFlow.")

    # 6. Version and Metadata
    add_heading(doc, "6. Version and Metadata", level=1)
    def extract_metadata_from_xml(root):
        metadata = {}
        for prop in root.findall(".//{http:///com.sap.ifl.model/Ifl.xsd}property"):
            key = prop.findtext("key")
            value = prop.findtext("value")
            if key and value:
                if key.lower() in ["componentversion", "author", "description", "componentns", "componentswcvname", "componentswcvid"]:
                    metadata[key] = value
        return metadata
    metadata = extract_metadata_from_xml(root)
    if metadata:
        add_colored_table(doc, [[k, v] for k, v in metadata.items()], ["Key", "Value"])
    else:
        add_paragraph(doc, "No metadata found in XML.")
    metadata_xml = ("<Metadata>\n" + "\n".join([f"<{k}>{v}</{k}>" for k, v in metadata.items()]) + "\n</Metadata>")
    metadata_summary = gemini_section_summary("Version and Metadata", metadata_xml, "Summarize the key metadata and versioning information for this SAP iFlow.")
    add_paragraph(doc, metadata_summary)

    # 7. Appendix
    add_heading(doc, "7. Appendix", level=1)
    def extract_appendix_info(root):
        appendix = []
        for prop in root.findall(".//{http:///com.sap.ifl.model/Ifl.xsd}property"):
            key = prop.findtext("key")
            value = prop.findtext("value")
            if key and value and key.lower().startswith("mapping"):
                appendix.append((key, value))
        return appendix
    appendix_xml = extract_section_xml(root, ".//{http://www.omg.org/spec/BPMN/20100524/MODEL}process")
    appendix_summary = gemini_section_summary("Appendix", appendix_xml, "List and briefly describe all technical artifacts, mappings, and scripts referenced in this iFlow.")
    add_paragraph(doc, appendix_summary)
    appendix_info = extract_appendix_info(root)
    if appendix_info:
        add_colored_table(doc, appendix_info, ["Key", "Value"])
    else:
        add_paragraph(doc, "No additional appendix info found in XML.")

    # Save locally (temp file) and then upload
    # Ensure file has a proper .docx name
    if not DOCX_PATH.lower().endswith(".docx"):
        DOCX_PATH_MOD = f"{DOCX_PATH}.docx"
    else:
        DOCX_PATH_MOD = DOCX_PATH

    # Save to current working directory first
    local_output_path = os.path.abspath(os.path.basename(DOCX_PATH_MOD))
    doc.save(iFlow_Documentation.docx)
    print(f"✅ Document generated and saved locally at: {local_output_path}")

    # Attempt to upload to GitHub if token + repo are provided
    if GITHUB_REPO and GITHUB_TOKEN:
        upload_result = upload_to_github(local_output_path, GITHUB_REPO, branch=GITHUB_BRANCH, repo_folder=REPO_TARGET_PATH, token=GITHUB_TOKEN)
        if upload_result.get("ok"):
            print("Upload successful:", upload_result.get("url"))
        else:
            print("Upload failed or skipped:", upload_result)
    else:
        print("GITHUB_REPO or GITHUB_TOKEN not configured. Document saved locally. To automatically push, configure 'github_repo' and 'github_token' in config_file.json.")

if __name__ == "__main__":
    generate_iflow_spec()
